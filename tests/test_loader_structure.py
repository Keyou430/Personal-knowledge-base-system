from docx import Document as DocxDocument

from core.loader import load_docx


def test_load_docx_preserves_heading_and_table(tmp_path):
    path = tmp_path / "制度.docx"
    doc = DocxDocument()
    doc.add_heading("报销流程", level=1)
    doc.add_paragraph("员工提交材料后进入审批。")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "角色"
    table.rows[0].cells[1].text = "审批人"
    doc.save(path)

    docs = load_docx(str(path))

    assert any(item.metadata["block_type"] == "heading" for item in docs)
    table_doc = next(item for item in docs if item.metadata["block_type"] == "table")
    assert "审批人" in table_doc.page_content
    assert table_doc.metadata["heading_path"] == "报销流程"
